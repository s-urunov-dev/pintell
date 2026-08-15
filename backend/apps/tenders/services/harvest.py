"""Mirror the documents a notice points at, before the links stop working.

A notice's real requirements are almost never in the notice. They are in the
Terms of Reference or the bidding document, and those live on the borrower's
own infrastructure — a ministry file server, a shared Drive folder — not on any
World Bank endpoint. Two consequences shape this module:

* **The links expire.** When a tender closes, the folder is taken down or made
  private. The notice keeps the URL forever; the file behind it is gone within
  months. So a document not mirrored today is one nobody can read next year,
  and harvesting runs from the first day rather than when something consumes it.
* **The hosts are not APIs.** These are somebody's web server with no published
  quota, so every default here is gentle: one request at a time per host, a
  pause between them, a hard size cap, and a long backoff on failure.

What comes back is stored twice: the raw bytes on disk, content-addressed by
SHA-256, and the extracted text in the database. The bytes are kept because
parsing improves — a scanned document that yields nothing today can be OCR'd
later without asking the borrower's server for it again.

The harvest also *measures* the corpus as a side effect. How many links are
public, how many sit behind a sign-in wall, how many are scans with no text
layer — those are product-defining numbers, and every fetch records its own
answer in ``status`` and ``has_text_layer``.
"""

from __future__ import annotations

import hashlib
import ipaddress
import logging
import re
import socket
import time
from dataclasses import dataclass, field
from datetime import timedelta
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, urlparse, urlunparse

import requests
from django.conf import settings
from django.utils import timezone
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry

from ..models import HarvestedDocument, ProjectDocument, ProjectProfile, TenderNotice
from ..notice_links import KIND_BIDDING, KIND_TOR, extract_links
from ..sanitizers import html_to_text

logger = logging.getLogger(__name__)

#: ``notice_links`` kinds → the storage vocabulary. Both modules classify the
#: same thing, but ``notice_links`` answers "what did the sentence call it"
#: while this one answers "what is it in the corpus", so the mapping is
#: explicit rather than the two sharing constants.
_KIND_MAP = {
    KIND_TOR: HarvestedDocument.Kind.TOR,
    KIND_BIDDING: HarvestedDocument.Kind.BIDDING,
}

# --- content sniffing ------------------------------------------------------
# File type is decided by the leading bytes, never by the URL suffix or the
# Content-Type header: borrowers serve PDFs as `application/octet-stream`, and
# a Drive link that has expired returns `text/html` with a `.pdf` in the path.
_MAGIC_PDF = b"%PDF"
_MAGIC_ZIP = b"PK\x03\x04"  # DOCX/XLSX are zip containers.

#: Markers of a page that is a door rather than a document. Checked only when
#: the response was HTML, so a PDF that happens to contain the word "sign in"
#: is unaffected.
_ACCESS_WALL_RE = re.compile(
    r"request access|sign in|log in to continue|you need permission|"
    r"accès refusé|acceso denegado|403 forbidden|not authorized",
    re.IGNORECASE,
)

#: Google Drive/Docs ids. These links dominate the TOR corpus and none of them
#: serve the file at the URL printed in the notice — see `_direct_url`.
_DRIVE_FILE_RE = re.compile(r"drive\.google\.com/file/d/([\w-]{10,})")
_DOCS_RE = re.compile(
    r"docs\.google\.com/(document|spreadsheets|presentation)/d/([\w-]{10,})"
)

_DOCS_EXPORT = {
    "document": "pdf",
    "spreadsheets": "xlsx",
    "presentation": "pdf",
}


class HarvestRejected(RuntimeError):
    """A URL that will not be fetched at all, with the reason why."""


@dataclass
class HarvestStats:
    discovered: int = 0
    attempted: int = 0
    fetched: int = 0
    access_denied: int = 0
    unreachable: int = 0
    no_text: int = 0
    skipped: int = 0
    #: Bytes that turned out to be a document we already had under another URL.
    deduplicated: int = 0
    errors: list[str] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "discovered": self.discovered,
            "attempted": self.attempted,
            "fetched": self.fetched,
            "access_denied": self.access_denied,
            "unreachable": self.unreachable,
            "no_text": self.no_text,
            "skipped": self.skipped,
            "deduplicated": self.deduplicated,
            "errors": self.errors[:20],
        }


# ---------------------------------------------------------------------------
# URL handling
# ---------------------------------------------------------------------------
def normalise_url(raw: str) -> str:
    """Canonical form of a URL, so one document is not fetched twice.

    Only the parts that cannot change what is served are touched: the scheme
    and host are lower-cased and the fragment is dropped. Query strings are
    left exactly as they are — on a Drive link the id lives there, and on a
    borrower's CMS ``?download=1`` is the difference between the file and a
    landing page.
    """
    parsed = urlparse((raw or "").strip())
    return urlunparse(
        (
            parsed.scheme.lower(),
            parsed.netloc.lower(),
            parsed.path,
            parsed.params,
            parsed.query,
            "",  # fragment
        )
    )


def url_key(url: str) -> str:
    return hashlib.sha256(normalise_url(url).encode("utf-8")).hexdigest()


def _direct_url(url: str) -> str:
    """Rewrite a share link into the URL that actually serves bytes.

    Google hosts most of the TOR corpus and none of its share links return a
    file: ``/file/d/<id>/view`` is a viewer page, and ``/document/d/<id>/edit``
    is an editor. Fetching either stores an HTML shell and reports the document
    as unreadable, which would be a wrong answer rather than a missing one.

    Anything not recognised is returned unchanged.
    """
    match = _DRIVE_FILE_RE.search(url)
    if match:
        return f"https://drive.google.com/uc?export=download&id={match.group(1)}"

    match = _DOCS_RE.search(url)
    if match:
        kind, doc_id = match.group(1), match.group(2)
        fmt = _DOCS_EXPORT.get(kind, "pdf")
        return f"https://docs.google.com/{kind}/d/{doc_id}/export?format={fmt}"

    # `open?id=` is the older share form for both products.
    parsed = urlparse(url)
    if parsed.netloc.endswith("google.com") and parsed.path == "/open":
        ids = parse_qs(parsed.query).get("id")
        if ids:
            return f"https://drive.google.com/uc?export=download&id={ids[0]}"

    return url


def _assert_fetchable(url: str) -> None:
    """Refuse URLs that must never be requested.

    These addresses come out of third-party notice text, so treating them as
    fetchable by default would let whoever writes a notice aim this service at
    anything the container can reach — the cloud metadata endpoint, the
    database, an internal admin port. The host is resolved and every address it
    answers with is checked, because a public name can resolve to a private
    address and only the resolved value tells the truth.
    """
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise HarvestRejected(f"unsupported scheme {parsed.scheme!r}")

    host = parsed.hostname
    if not host:
        raise HarvestRejected("no host in URL")

    try:
        resolved = socket.getaddrinfo(host, None)
    except socket.gaierror as exc:
        raise HarvestRejected(f"host does not resolve: {exc}") from exc

    for entry in resolved:
        address = ipaddress.ip_address(entry[4][0])
        if (
            address.is_private
            or address.is_loopback
            or address.is_link_local
            or address.is_reserved
            or address.is_multicast
        ):
            raise HarvestRejected(f"host resolves to a non-public address ({address})")


# ---------------------------------------------------------------------------
# Discovery — which links exist at all
# ---------------------------------------------------------------------------
def discover_from_notices(
    queryset=None, *, limit: int = 500, stats: HarvestStats | None = None
) -> HarvestStats:
    """Read notice bodies and register every document link they carry.

    Registering is not fetching: this only writes rows in ``pending``, so the
    expensive half stays bounded by its own batch size and schedule. A link
    already known simply gains another notice in ``notices``.
    """
    stats = stats or HarvestStats()
    if queryset is None:
        queryset = TenderNotice.objects.focus()

    notices = queryset.exclude(notice_text_sanitized="").only(
        "notice_id", "notice_text_sanitized"
    )[:limit]

    for notice in notices:
        text = html_to_text(notice.notice_text_sanitized)
        for link in extract_links(text):
            document, created = _register(
                link.url,
                kind=_KIND_MAP.get(link.kind, HarvestedDocument.Kind.OTHER),
                context=link.context,
            )
            document.notices.add(notice)
            if created:
                stats.discovered += 1

    return stats


def discover_from_project_documents(
    *, limit: int = 500, stats: HarvestStats | None = None
) -> HarvestStats:
    """Register the project PDFs already mirrored as metadata.

    ``ProjectDocument`` stores a title and a link; the file itself was never
    fetched. Those are first-party World Bank URLs and therefore the stable
    part of the corpus — worth harvesting precisely because the borrower-hosted
    half is not.

    **Only rows not registered yet, and that is a fix rather than an
    optimisation.** This used to slice ``limit`` rows off an unordered
    queryset every cycle, which is not a page — Postgres is free to return the
    same rows each time, and it did. Measured on the deployed server on
    2026-08-14: 1,559 project documents, 711 ever registered, and the remaining
    848 unreachable by any number of further runs. Among them were 17 of the 19
    ESRS files this app now reads its Bank-side contacts out of, which is how
    the gap was found at all. Excluding what is already known makes each run
    advance, and makes an exhausted corpus cost one query.

    ``ProjectProfile.esrs_pdf_url`` is registered alongside them: the ESRS is a
    project document like the rest, but it is the one this codebase asks for by
    name (:mod:`apps.tenders.esrs`), so it should not depend on arriving in the
    right slice of a bounded batch.
    """
    stats = stats or HarvestStats()

    # Every URL already registered, keyed the way ``_register`` keys them. Held
    # in memory rather than filtered in SQL because ``ProjectDocument`` stores
    # the raw URL and ``HarvestedDocument`` the normalised hash of it, so there
    # is no column to join on — and because it is what makes the scan below
    # cost one row read per candidate instead of one query.
    known = set(HarvestedDocument.objects.values_list("url_hash", flat=True))
    registered = 0

    def register(url: str, title: str) -> bool:
        """Register one URL. False when it was already known, so it is not a slot."""
        nonlocal registered
        key = url_key(url) if url else ""
        if not key or key in known:
            return False
        known.add(key)
        registered += 1
        _, created = _register(
            url, kind=HarvestedDocument.Kind.PROJECT_DOC, context=title or ""
        )
        if created:
            stats.discovered += 1
        return True

    # The ESRS first, so a full batch still reaches the file a reader is waiting
    # on rather than the twelfth revision of a procurement plan.
    sources = [
        ProjectProfile.objects.exclude(esrs_pdf_url="")
        .order_by("project_id")
        .values_list("esrs_pdf_url", "esrs_title"),
        ProjectDocument.objects.exclude(pdf_url="")
        .order_by("-doc_date", "guid")
        .values_list("pdf_url", "title"),
    ]
    # The limit counts *new* rows, and the scan walks past the known ones to
    # find them. Slicing the queryset instead is what broke this: a slice of an
    # ordered list of candidates is the same slice every cycle once its rows are
    # registered, so the batch reported a full run and discovered nothing.
    for source in sources:
        for url, title in source.iterator():
            register(url, title)
            if registered >= limit:
                return stats
    return stats


def _register(url: str, *, kind: str, context: str) -> tuple[HarvestedDocument, bool]:
    """Create or update the row for one URL, keeping the most specific kind."""
    key = url_key(url)
    document, created = HarvestedDocument.objects.get_or_create(
        url_hash=key,
        defaults={
            "url": normalise_url(url),
            "kind": kind,
            "link_context": context[:500],
        },
    )
    # The same URL can be introduced by a describing sentence in one notice and
    # bare in another; the specific reading wins, exactly as in `notice_links`.
    if not created and document.kind == HarvestedDocument.Kind.OTHER and kind != document.kind:
        document.kind = kind
        document.link_context = context[:500] or document.link_context
        document.save(update_fields=["kind", "link_context", "updated_at"])
    return document, created


# ---------------------------------------------------------------------------
# Fetching
# ---------------------------------------------------------------------------
def _build_session() -> requests.Session:
    session = requests.Session()
    retry = Retry(
        total=2, connect=2, read=2, backoff_factor=1.5,
        status_forcelist=(429, 500, 502, 503, 504),
        allowed_methods=frozenset({"GET"}), raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retry, pool_maxsize=4)
    session.mount("https://", adapter)
    session.mount("http://", adapter)
    session.headers.update(
        {
            "User-Agent": settings.WORLDBANK["USER_AGENT"],
            # Deliberately broad: borrowers serve documents under a wide range
            # of content types, and a narrow Accept gets a 406 from some CMSs.
            "Accept": "*/*",
        }
    )
    return session


def _download(session: requests.Session, url: str) -> tuple[requests.Response, bytes]:
    """Stream a response, stopping at the size cap.

    Streaming rather than ``response.content`` because the cap has to apply to
    what is *read*, not to what was already pulled into memory — a mis-linked
    archive would otherwise be downloaded in full before being rejected.
    """
    config = settings.HARVEST
    response = session.get(
        url, timeout=config["HTTP_TIMEOUT"], stream=True, allow_redirects=True
    )

    chunks: list[bytes] = []
    total = 0
    for chunk in response.iter_content(chunk_size=64 * 1024):
        if not chunk:
            continue
        chunks.append(chunk)
        total += len(chunk)
        if total > config["MAX_BYTES"]:
            response.close()
            raise HarvestRejected(f"exceeds size cap ({config['MAX_BYTES']} bytes)")
    response.close()
    return response, b"".join(chunks)


def _store(payload: bytes, suffix: str) -> tuple[str, str]:
    """Write bytes to content-addressed storage. Returns ``(sha256, path)``.

    Two URLs serving an identical file — common when a borrower re-uploads an
    unchanged document — collapse onto one blob, and a re-fetch of the same
    document overwrites itself byte for byte.
    """
    digest = hashlib.sha256(payload).hexdigest()
    root: Path = settings.HARVEST["DIR"]
    target = root / digest[:2] / f"{digest}{suffix}"
    target.parent.mkdir(parents=True, exist_ok=True)
    if not target.exists():
        target.write_bytes(payload)
    return digest, str(target)


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text(payload: bytes, content_type: str) -> dict[str, Any]:
    """Pull plain text out of whatever came back.

    Returns the parser used, the text, the page count where the format has one,
    and ``has_text_layer`` — the field that separates "we cannot read this" from
    "there is nothing to read", which is the difference between a bug and a scan.
    """
    if payload[:4] == _MAGIC_PDF:
        return _extract_pdf(payload)
    if payload[:4] == _MAGIC_ZIP:
        return _extract_docx(payload)

    # Neither magic matched. If it looks like markup, flatten it — some
    # borrowers publish the TOR as a web page rather than a file.
    head = payload[:2048].lower()
    if b"<html" in head or b"<!doctype html" in head or "html" in content_type:
        text = html_to_text(payload.decode("utf-8", errors="replace"))
        return {
            "parser": "html",
            "text": text,
            "page_count": None,
            "has_text_layer": bool(text.strip()),
        }

    # Plain text, but only when the server (or, for an upload, the file
    # extension) actually says so. Sniffing instead would decode any
    # unrecognised binary into replacement characters and report a document
    # full of text that says nothing — worse than reporting we cannot read it,
    # because the first is a wrong answer and the second is a missing one.
    if content_type.split(";")[0].strip().startswith("text/"):
        text = payload.decode("utf-8", errors="replace")
        return {
            "parser": "text",
            "text": text,
            "page_count": None,
            "has_text_layer": bool(text.strip()),
        }

    return {
        "parser": "",
        "text": "",
        "page_count": None,
        "has_text_layer": None,
        "error": f"unrecognised format (content-type {content_type or 'unknown'})",
    }


def _extract_pdf(payload: bytes) -> dict[str, Any]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return {
            "parser": "", "text": "", "page_count": None, "has_text_layer": None,
            "error": "pypdf is not installed",
        }

    import io

    try:
        reader = PdfReader(io.BytesIO(payload))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except Exception as exc:  # noqa: BLE001 - a malformed PDF is data, not a bug
        return {
            "parser": "pypdf", "text": "", "page_count": None, "has_text_layer": None,
            "error": f"pypdf failed: {exc}",
        }

    text = "\n\n".join(pages).strip()
    return {
        "parser": "pypdf",
        "text": text,
        "page_count": len(pages),
        # A PDF that parsed cleanly but yielded almost nothing is a scan. Saying
        # so explicitly is what tells us whether OCR is worth building at all.
        "has_text_layer": len(text) >= HarvestedDocument.MIN_USEFUL_CHARS,
    }


def _extract_docx(payload: bytes) -> dict[str, Any]:
    try:
        import docx  # python-docx
    except ImportError:
        return {
            "parser": "", "text": "", "page_count": None, "has_text_layer": None,
            "error": "python-docx is not installed",
        }

    import io

    try:
        document = docx.Document(io.BytesIO(payload))
    except Exception as exc:  # noqa: BLE001 - also covers "this zip is an xlsx"
        return {
            "parser": "docx", "text": "", "page_count": None, "has_text_layer": None,
            "error": f"python-docx failed: {exc}",
        }

    parts = [p.text for p in document.paragraphs]
    # Qualification criteria live in tables far more often than in paragraphs,
    # so the cells are read too — dropping them would lose exactly the section
    # this whole pipeline exists to reach.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    text = "\n".join(part for part in parts if part.strip()).strip()
    return {
        "parser": "docx",
        "text": text,
        "page_count": None,
        "has_text_layer": bool(text),
    }


def _suffix_for(parser: str, payload: bytes) -> str:
    if payload[:4] == _MAGIC_PDF:
        return ".pdf"
    if payload[:4] == _MAGIC_ZIP:
        return ".docx"
    return ".html" if parser == "html" else ".bin"


# ---------------------------------------------------------------------------
# One document
# ---------------------------------------------------------------------------
def _retry_delay(attempts: int) -> timedelta:
    config = settings.HARVEST
    exponent = min(max(attempts - 1, 0), 16)
    minutes = config["RETRY_BASE_MINUTES"] * (2**exponent)
    return min(timedelta(minutes=minutes), timedelta(days=config["RETRY_MAX_DAYS"]))


def _finish(
    document: HarvestedDocument, status: str, *, error: str = "", now=None
) -> HarvestedDocument:
    """Record the outcome of one attempt, with backoff when it can be retried."""
    now = now or timezone.now()
    document.status = status
    document.last_attempt_at = now
    document.last_error = error[:2000]

    if status == HarvestedDocument.Status.FETCHED:
        document.fetched_at = now
        document.next_retry_at = None
    elif status in HarvestedDocument.RETRYABLE_STATUSES:
        document.next_retry_at = now + _retry_delay(document.attempts)
    else:
        # NO_TEXT and SKIPPED are settled answers: the bytes are on disk or the
        # URL was refused outright, and asking again changes neither.
        document.next_retry_at = None

    document.save()
    return document


def harvest_document(
    document: HarvestedDocument,
    *,
    session: requests.Session | None = None,
    stats: HarvestStats | None = None,
) -> HarvestedDocument:
    """Fetch, store and parse one document. Never raises."""
    stats = stats or HarvestStats()
    session = session or _build_session()
    config = settings.HARVEST

    document.attempts += 1
    stats.attempted += 1

    target = _direct_url(document.url)
    try:
        _assert_fetchable(target)
    except HarvestRejected as exc:
        stats.skipped += 1
        return _finish(document, HarvestedDocument.Status.SKIPPED, error=str(exc))

    try:
        response, payload = _download(session, target)
    except HarvestRejected as exc:
        stats.skipped += 1
        return _finish(document, HarvestedDocument.Status.SKIPPED, error=str(exc))
    except requests.RequestException as exc:
        stats.unreachable += 1
        return _finish(document, HarvestedDocument.Status.UNREACHABLE, error=str(exc))

    document.http_status = response.status_code
    document.content_type = (response.headers.get("Content-Type") or "")[:128]
    document.byte_size = len(payload)

    if response.status_code >= 400:
        stats.unreachable += 1
        return _finish(
            document,
            HarvestedDocument.Status.UNREACHABLE,
            error=f"HTTP {response.status_code}",
        )

    if not payload:
        stats.unreachable += 1
        return _finish(document, HarvestedDocument.Status.UNREACHABLE, error="empty body")

    # An HTML body where a document was expected is the signature of a sign-in
    # wall — Drive answers 200 with its "request access" page. Recording that
    # as its own status is what makes "how much of the corpus is actually
    # public" a measurable number rather than a guess.
    looks_like_page = payload[:4] not in (_MAGIC_PDF, _MAGIC_ZIP)
    if looks_like_page and _ACCESS_WALL_RE.search(payload[:20_000].decode("utf-8", "replace")):
        stats.access_denied += 1
        return _finish(
            document,
            HarvestedDocument.Status.ACCESS_DENIED,
            error="response is an access wall, not a document",
        )

    result = extract_text(payload, document.content_type)
    document.parser = result["parser"]
    document.page_count = result["page_count"]
    document.has_text_layer = result["has_text_layer"]
    document.parse_error = str(result.get("error", ""))[:1000]

    # NUL bytes are stripped before the truncation, and this is not defensive
    # tidying — Postgres refuses a text field containing 0x00 outright. Found on
    # 2026-08-14, the first time the harvester reached the project documents an
    # unordered-slice bug had been hiding from it (D66): one PDF parses to text
    # with embedded NULs, and it took a batch of 120 down with it *despite* the
    # per-document guard in `harvest_pending`. That is the part worth
    # understanding — the DataError surfaces from the save, so the guard caught
    # it and then called `_finish` to record the failure, which saved the same
    # instance with the same NULs and raised again, this time outside the try.
    # A parser artefact must produce a recorded state, like every other
    # malformed document here.
    text = (result["text"] or "").replace("\x00", "")[: config["MAX_TEXT_CHARS"]]
    document.text = text
    document.text_chars = len(text)

    # Bytes are stored even when nothing could be read out of them: the file is
    # the asset, and a better parser (or OCR) should not have to re-fetch it
    # from a link that may not answer twice.
    existing = (
        HarvestedDocument.objects.filter(sha256=hashlib.sha256(payload).hexdigest())
        .exclude(pk=document.pk)
        .exists()
    )
    digest, path = _store(payload, _suffix_for(document.parser, payload))
    document.sha256 = digest
    document.stored_path = path
    if existing:
        stats.deduplicated += 1

    if document.text_chars < HarvestedDocument.MIN_USEFUL_CHARS:
        stats.no_text += 1
        return _finish(
            document,
            HarvestedDocument.Status.NO_TEXT,
            error=document.parse_error or "no usable text layer",
        )

    stats.fetched += 1
    return _finish(document, HarvestedDocument.Status.FETCHED)


# ---------------------------------------------------------------------------
# Batch
# ---------------------------------------------------------------------------
def select_pending(limit: int) -> list[HarvestedDocument]:
    """Never-tried documents first, then failures whose backoff has elapsed.

    A link nobody has tried outranks a retry for the same reason a new project
    outranks a stale one in ``services.projects``: the unknown is where the
    information is.
    """
    fresh = list(HarvestedDocument.objects.never_tried().order_by("created_at")[:limit])
    remaining = limit - len(fresh)
    if remaining <= 0:
        return fresh
    return fresh + list(
        HarvestedDocument.objects.retry_due().order_by("next_retry_at")[:remaining]
    )


def harvest_pending(*, limit: int | None = None) -> HarvestStats:
    """Fetch one bounded batch, pausing between requests to the same host."""
    stats = HarvestStats()
    if not settings.HARVEST["ENABLED"]:
        logger.info("Harvest disabled (HARVEST_ENABLED=false) — skipping.")
        return stats

    limit = limit or settings.HARVEST["BATCH_SIZE"]
    documents = select_pending(limit)
    if not documents:
        return stats

    session = _build_session()
    delay = settings.HARVEST["PER_HOST_DELAY"]
    last_seen_host: dict[str, float] = {}

    for document in documents:
        host = urlparse(document.url).netloc
        # These are somebody's file server, not an API with a quota. Spacing
        # requests per host rather than globally keeps a batch spread across
        # twenty borrowers fast while still being gentle on each of them.
        previous = last_seen_host.get(host)
        if previous is not None:
            elapsed = time.monotonic() - previous
            if elapsed < delay:
                time.sleep(delay - elapsed)

        try:
            harvest_document(document, session=session, stats=stats)
        except Exception as exc:  # noqa: BLE001 - one link must not kill a batch
            logger.exception("Harvest failed for %s", document.url[:120])
            stats.errors.append(f"{document.url[:80]}: {exc}")
            # The half-written attempt is discarded before the failure is
            # recorded. Recording it means saving this instance, and the field
            # that made the attempt fail is very often the field that would
            # make the recording fail too — which is how one document took a
            # whole batch down rather than itself. There is nothing to lose:
            # an attempt that raised produced no text worth keeping.
            document.text = ""
            document.text_chars = 0
            _finish(document, HarvestedDocument.Status.UNREACHABLE, error=str(exc))

        last_seen_host[host] = time.monotonic()

    logger.info("Harvest batch complete: %s", stats.as_dict())
    return stats


def corpus_report() -> dict[str, Any]:
    """What the harvest has actually reached, as counts.

    This answers the questions the product depends on — how much of the linked
    corpus is public, how much is a scan — from the whole corpus rather than
    from a hand-checked sample of ten.
    """
    from django.db.models import Avg, Count, Q

    totals = HarvestedDocument.objects.aggregate(
        total=Count("url_hash"),
        pending=Count("url_hash", filter=Q(status=HarvestedDocument.Status.PENDING)),
        fetched=Count("url_hash", filter=Q(status=HarvestedDocument.Status.FETCHED)),
        access_denied=Count(
            "url_hash", filter=Q(status=HarvestedDocument.Status.ACCESS_DENIED)
        ),
        unreachable=Count(
            "url_hash", filter=Q(status=HarvestedDocument.Status.UNREACHABLE)
        ),
        no_text=Count("url_hash", filter=Q(status=HarvestedDocument.Status.NO_TEXT)),
        skipped=Count("url_hash", filter=Q(status=HarvestedDocument.Status.SKIPPED)),
        scanned=Count("url_hash", filter=Q(has_text_layer=False)),
        avg_chars=Avg("text_chars", filter=Q(status=HarvestedDocument.Status.FETCHED)),
    )

    by_kind = {
        row["kind"]: row["count"]
        for row in HarvestedDocument.objects.values("kind").annotate(
            count=Count("url_hash")
        )
    }

    attempted = totals["total"] - totals["pending"]
    return {
        **totals,
        "by_kind": by_kind,
        "attempted": attempted,
        # The single number this whole module exists to establish: of the links
        # we tried, how many produced a readable document.
        "reachable_rate": (
            round(totals["fetched"] / attempted, 3) if attempted else None
        ),
    }
