"""Tests for the linked-document harvester.

Nothing here touches the network: the fetch path takes a session object, so the
tests hand it a stub that returns whatever byte string the case is about.
"""

from __future__ import annotations

import hashlib
import io
from unittest import mock

from django.test import TestCase, override_settings
from django.utils import timezone

from apps.tenders.models import HarvestedDocument, TenderNotice
from apps.tenders.services import harvest


class FakeResponse:
    """Enough of ``requests.Response`` for :func:`harvest._download`."""

    def __init__(self, payload: bytes, *, status: int = 200, content_type: str = ""):
        self._payload = payload
        self.status_code = status
        self.headers = {"Content-Type": content_type}

    def iter_content(self, chunk_size: int = 0):
        yield self._payload

    def close(self) -> None:
        pass


class FakeSession:
    def __init__(self, response: FakeResponse):
        self._response = response
        self.requested: list[str] = []

    def get(self, url, **kwargs):
        self.requested.append(url)
        return self._response


def _docx_bytes(paragraph: str = "", table_rows: list[list[str]] | None = None) -> bytes:
    import docx

    document = docx.Document()
    if paragraph:
        document.add_paragraph(paragraph)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, value in enumerate(row):
                table.cell(row_index, col_index).text = value
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _blank_pdf_bytes(pages: int = 1) -> bytes:
    from pypdf import PdfWriter

    writer = PdfWriter()
    for _ in range(pages):
        writer.add_blank_page(width=200, height=200)
    buffer = io.BytesIO()
    writer.write(buffer)
    return buffer.getvalue()


class UrlNormalisationTests(TestCase):
    def test_scheme_and_host_are_folded_but_the_query_is_not(self):
        """The query string decides what is served, so it survives untouched."""
        normalised = harvest.normalise_url("HTTPS://Docs.Example.COM/a?B=1&a=2#frag")
        self.assertEqual(normalised, "https://docs.example.com/a?B=1&a=2")

    def test_two_spellings_of_one_url_share_a_key(self):
        self.assertEqual(
            harvest.url_key("https://example.com/tor.pdf#page=3"),
            harvest.url_key("https://EXAMPLE.com/tor.pdf"),
        )

    def test_differing_queries_stay_separate(self):
        self.assertNotEqual(
            harvest.url_key("https://x.io/get?id=1"),
            harvest.url_key("https://x.io/get?id=2"),
        )


class DirectUrlTests(TestCase):
    """Google share links never serve the file at the printed URL."""

    def test_drive_file_link_becomes_a_download(self):
        self.assertEqual(
            harvest._direct_url(
                "https://drive.google.com/file/d/1AbCdEfGhIjKlMnO/view?usp=sharing"
            ),
            "https://drive.google.com/uc?export=download&id=1AbCdEfGhIjKlMnO",
        )

    def test_docs_link_becomes_a_pdf_export(self):
        self.assertEqual(
            harvest._direct_url(
                "https://docs.google.com/document/d/1AbCdEfGhIjKlMnO/edit?usp=sharing"
            ),
            "https://docs.google.com/document/d/1AbCdEfGhIjKlMnO/export?format=pdf",
        )

    def test_spreadsheet_exports_as_a_spreadsheet(self):
        self.assertIn(
            "format=xlsx",
            harvest._direct_url("https://docs.google.com/spreadsheets/d/1AbCdEfGhIjKl/edit"),
        )

    def test_legacy_open_id_form(self):
        self.assertEqual(
            harvest._direct_url("https://drive.google.com/open?id=1AbCdEfGhIjKlMnO"),
            "https://drive.google.com/uc?export=download&id=1AbCdEfGhIjKlMnO",
        )

    def test_an_ordinary_url_is_left_alone(self):
        url = "https://ministry.gov.uz/files/tor.pdf"
        self.assertEqual(harvest._direct_url(url), url)


class FetchGuardTests(TestCase):
    """These URLs come out of third-party notice text and are not trusted."""

    def test_non_http_scheme_is_refused(self):
        with self.assertRaises(harvest.HarvestRejected):
            harvest._assert_fetchable("file:///etc/passwd")

    def test_private_address_is_refused(self):
        with mock.patch.object(
            harvest.socket, "getaddrinfo",
            return_value=[(0, 0, 0, "", ("169.254.169.254", 0))],
        ):
            with self.assertRaises(harvest.HarvestRejected):
                harvest._assert_fetchable("https://metadata.example.com/latest")

    def test_loopback_is_refused_even_behind_a_public_name(self):
        with mock.patch.object(
            harvest.socket, "getaddrinfo",
            return_value=[(0, 0, 0, "", ("127.0.0.1", 0))],
        ):
            with self.assertRaises(harvest.HarvestRejected):
                harvest._assert_fetchable("https://looks-public.example.com/x")

    def test_public_address_passes(self):
        with mock.patch.object(
            harvest.socket, "getaddrinfo",
            return_value=[(0, 0, 0, "", ("93.184.216.34", 0))],
        ):
            harvest._assert_fetchable("https://example.com/tor.pdf")


class TextExtractionTests(TestCase):
    def test_format_is_decided_by_magic_bytes_not_by_the_header(self):
        """Borrowers mislabel content types constantly; the bytes do not lie."""
        result = harvest.extract_text(_blank_pdf_bytes(), "application/octet-stream")
        self.assertEqual(result["parser"], "pypdf")

    def test_a_pdf_without_a_text_layer_reports_the_fact(self):
        result = harvest.extract_text(_blank_pdf_bytes(pages=2), "application/pdf")
        self.assertEqual(result["page_count"], 2)
        self.assertFalse(result["has_text_layer"])

    def test_docx_paragraphs_are_read(self):
        payload = _docx_bytes(paragraph="Minimum average annual turnover of US$10,000,000")
        result = harvest.extract_text(payload, "")
        self.assertEqual(result["parser"], "docx")
        self.assertIn("US$10,000,000", result["text"])

    def test_docx_tables_are_read(self):
        """Qualification criteria live in tables far more often than in prose."""
        payload = _docx_bytes(
            table_rows=[
                ["Criterion", "Single Entity", "Each Party"],
                ["Similar contracts", "2", "1"],
            ]
        )
        result = harvest.extract_text(payload, "")
        self.assertIn("Similar contracts", result["text"])
        self.assertIn("Single Entity", result["text"])

    def test_html_is_flattened(self):
        payload = b"<html><body><p>Terms of Reference</p></body></html>"
        result = harvest.extract_text(payload, "text/html")
        self.assertEqual(result["parser"], "html")
        self.assertIn("Terms of Reference", result["text"])

    def test_an_unrecognised_format_says_so_rather_than_guessing(self):
        result = harvest.extract_text(b"\x00\x01\x02binary", "application/x-thing")
        self.assertEqual(result["parser"], "")
        self.assertIn("unrecognised", result["error"])


class DiscoveryTests(TestCase):
    def setUp(self):
        self.notice = TenderNotice.objects.create(
            notice_id="OP-DISCOVER-1",
            notice_type="Request for Expression of Interest",
            country="Uzbekistan",
            deadline_date=timezone.now() + timezone.timedelta(days=10),
            notice_text_sanitized=(
                "<p>The detailed Terms of Reference (TOR) for the assignment can be "
                "found at the following link: https://example.com/tor.pdf</p>"
                # `notice_links` reads 180 characters back for the phrase that
                # introduced a URL, so a second link needs real distance from
                # the first to be judged on its own context — which is how they
                # sit in an actual notice, paragraphs apart.
                "<p>Expressions of interest must be delivered in writing to the "
                "address below during office hours, that is from 09:00 to 17:00 "
                "local time, on or before the submission deadline stated above. "
                "Late submissions will be rejected.</p>"
                "<p>The agency home page is https://example.com/other.html</p>"
            ),
        )

    def test_links_are_registered_with_the_kind_the_sentence_gave_them(self):
        stats = harvest.discover_from_notices(
            queryset=TenderNotice.objects.filter(pk=self.notice.pk)
        )
        self.assertEqual(stats.discovered, 2)

        tor = HarvestedDocument.objects.get(url="https://example.com/tor.pdf")
        self.assertEqual(tor.kind, HarvestedDocument.Kind.TOR)
        self.assertEqual(tor.status, HarvestedDocument.Status.PENDING)
        self.assertIn("Terms of Reference", tor.link_context)

        other = HarvestedDocument.objects.get(url="https://example.com/other.html")
        self.assertEqual(other.kind, HarvestedDocument.Kind.OTHER)

    def test_the_notice_that_pointed_at_a_link_is_recorded(self):
        harvest.discover_from_notices(
            queryset=TenderNotice.objects.filter(pk=self.notice.pk)
        )
        tor = HarvestedDocument.objects.get(url="https://example.com/tor.pdf")
        self.assertEqual(list(tor.notices.values_list("pk", flat=True)), [self.notice.pk])

    def test_a_second_notice_pointing_at_one_document_does_not_duplicate_it(self):
        """One TOR shared by two notices is one row, and one future fetch."""
        second = TenderNotice.objects.create(
            notice_id="OP-DISCOVER-2",
            notice_type="Request for Expression of Interest",
            country="Uzbekistan",
            deadline_date=timezone.now() + timezone.timedelta(days=10),
            notice_text_sanitized=(
                "<p>Terms of Reference: https://example.com/tor.pdf</p>"
            ),
        )
        harvest.discover_from_notices(queryset=TenderNotice.objects.all())

        tor = HarvestedDocument.objects.get(url="https://example.com/tor.pdf")
        self.assertEqual(HarvestedDocument.objects.count(), 2)  # tor + other.html
        self.assertEqual(tor.notices.count(), 2)
        self.assertIn(second.pk, tor.notices.values_list("pk", flat=True))

    def test_rerunning_discovery_registers_nothing_new(self):
        queryset = TenderNotice.objects.filter(pk=self.notice.pk)
        harvest.discover_from_notices(queryset=queryset)
        again = harvest.discover_from_notices(queryset=queryset)
        self.assertEqual(again.discovered, 0)


@override_settings()
class HarvestDocumentTests(TestCase):
    """The fetch path, with the network stubbed out."""

    def setUp(self):
        self.tmp = self.enterContext(
            __import__("tempfile").TemporaryDirectory()
        )
        harvest_settings = dict(__import__("django.conf", fromlist=["settings"]).settings.HARVEST)
        harvest_settings["DIR"] = __import__("pathlib").Path(self.tmp)
        self.override = override_settings(HARVEST=harvest_settings)
        self.override.enable()
        self.addCleanup(self.override.disable)

        self.document = HarvestedDocument.objects.create(
            url_hash=harvest.url_key("https://example.com/tor.pdf"),
            url="https://example.com/tor.pdf",
            kind=HarvestedDocument.Kind.TOR,
        )
        patcher = mock.patch.object(harvest, "_assert_fetchable", return_value=None)
        patcher.start()
        self.addCleanup(patcher.stop)

    def _run(self, payload: bytes, *, status: int = 200, content_type: str = ""):
        session = FakeSession(FakeResponse(payload, status=status, content_type=content_type))
        return harvest.harvest_document(self.document, session=session), session

    def test_a_readable_document_is_stored_and_parsed(self):
        payload = _docx_bytes(paragraph="Minimum average annual turnover " + "x" * 300)
        document, _ = self._run(payload)

        self.assertEqual(document.status, HarvestedDocument.Status.FETCHED)
        self.assertEqual(document.parser, "docx")
        self.assertGreater(document.text_chars, HarvestedDocument.MIN_USEFUL_CHARS)
        self.assertEqual(document.sha256, hashlib.sha256(payload).hexdigest())
        self.assertTrue(__import__("pathlib").Path(document.stored_path).exists())
        self.assertIsNotNone(document.fetched_at)
        self.assertIsNone(document.next_retry_at)

    def test_an_access_wall_is_its_own_outcome(self):
        """Drive answers 200 with a sign-in page; that is not "unreachable"."""
        document, _ = self._run(
            b"<html><body>You need permission to access this document</body></html>",
            content_type="text/html",
        )
        self.assertEqual(document.status, HarvestedDocument.Status.ACCESS_DENIED)
        # Retried later: borrowers often open the folder up a few days on.
        self.assertIsNotNone(document.next_retry_at)

    def test_a_scan_is_stored_even_though_nothing_could_be_read(self):
        """The bytes are the asset — OCR must not have to re-fetch them."""
        document, _ = self._run(_blank_pdf_bytes())
        self.assertEqual(document.status, HarvestedDocument.Status.NO_TEXT)
        self.assertFalse(document.has_text_layer)
        self.assertTrue(document.stored_path)
        self.assertTrue(__import__("pathlib").Path(document.stored_path).exists())
        # A settled answer: asking again would return the same scan.
        self.assertIsNone(document.next_retry_at)

    def test_an_http_error_is_retryable(self):
        document, _ = self._run(b"nope", status=503)
        self.assertEqual(document.status, HarvestedDocument.Status.UNREACHABLE)
        self.assertEqual(document.attempts, 1)
        self.assertIsNotNone(document.next_retry_at)

    def test_a_parse_that_yields_nul_bytes_is_stored_rather_than_refused(self):
        """Postgres rejects 0x00 in a text field, from the save, killing the batch.

        Found when discovery first reached the project documents an unordered
        slice had been hiding (D66): one PDF parses to text with embedded NULs
        and took a batch of 120 down. The per-document guard did not help — it
        caught the DataError and then saved the same instance to record the
        failure, raising again outside the try.
        """
        text = "Minimum average annual turnover\x00 " + "x" * 300
        with mock.patch.object(
            harvest,
            "extract_text",
            return_value={
                "text": text, "parser": "pypdf", "page_count": 1,
                "has_text_layer": True, "error": "",
            },
        ):
            document, _ = self._run(_docx_bytes(paragraph="ignored"))

        self.assertEqual(document.status, HarvestedDocument.Status.FETCHED)
        self.assertNotIn("\x00", document.text)
        # Only the NULs go; the document is otherwise readable and countable.
        self.assertEqual(document.text_chars, len(text) - 1)
        document.refresh_from_db()
        self.assertNotIn("\x00", document.text)

    def test_the_share_link_is_rewritten_before_the_request(self):
        self.document.url = "https://drive.google.com/file/d/1AbCdEfGhIjKlMnO/view"
        self.document.save()
        _, session = self._run(_docx_bytes(paragraph="y" * 400))
        self.assertEqual(
            session.requested,
            ["https://drive.google.com/uc?export=download&id=1AbCdEfGhIjKlMnO"],
        )

    def test_a_refused_url_is_not_fetched_at_all(self):
        with mock.patch.object(
            harvest, "_assert_fetchable",
            side_effect=harvest.HarvestRejected("host resolves to a non-public address"),
        ):
            session = FakeSession(FakeResponse(b"secret"))
            document = harvest.harvest_document(self.document, session=session)

        self.assertEqual(document.status, HarvestedDocument.Status.SKIPPED)
        self.assertEqual(session.requested, [])
        self.assertIsNone(document.next_retry_at)

    def test_two_urls_serving_one_file_share_a_blob(self):
        payload = _docx_bytes(paragraph="z" * 400)
        first, _ = self._run(payload)

        twin = HarvestedDocument.objects.create(
            url_hash=harvest.url_key("https://mirror.example.com/tor.docx"),
            url="https://mirror.example.com/tor.docx",
            kind=HarvestedDocument.Kind.TOR,
        )
        session = FakeSession(FakeResponse(payload))
        second = harvest.harvest_document(twin, session=session)

        self.assertEqual(first.sha256, second.sha256)
        self.assertEqual(first.stored_path, second.stored_path)


class BackoffTests(TestCase):
    def test_the_delay_doubles_and_then_stops_growing(self):
        first = harvest._retry_delay(1)
        second = harvest._retry_delay(2)
        self.assertEqual(second, first * 2)

        capped = harvest._retry_delay(40)
        self.assertLessEqual(capped.days, 3)


class SelectionTests(TestCase):
    def test_never_tried_documents_come_before_retries(self):
        """The unknown is where the information is — same rule as projects."""
        overdue = HarvestedDocument.objects.create(
            url_hash="a" * 64, url="https://x.io/a", attempts=2,
            status=HarvestedDocument.Status.UNREACHABLE,
            next_retry_at=timezone.now() - timezone.timedelta(hours=1),
        )
        fresh = HarvestedDocument.objects.create(url_hash="b" * 64, url="https://x.io/b")

        selected = harvest.select_pending(limit=10)
        self.assertEqual([d.pk for d in selected], [fresh.pk, overdue.pk])

    def test_a_retry_that_is_not_due_is_left_alone(self):
        HarvestedDocument.objects.create(
            url_hash="c" * 64, url="https://x.io/c", attempts=1,
            status=HarvestedDocument.Status.UNREACHABLE,
            next_retry_at=timezone.now() + timezone.timedelta(hours=5),
        )
        self.assertEqual(harvest.select_pending(limit=10), [])

    def test_a_settled_failure_is_never_reselected(self):
        HarvestedDocument.objects.create(
            url_hash="d" * 64, url="https://x.io/d", attempts=1,
            status=HarvestedDocument.Status.NO_TEXT,
            next_retry_at=None,
        )
        self.assertEqual(harvest.select_pending(limit=10), [])


class CorpusReportTests(TestCase):
    def test_the_reachable_rate_counts_only_what_was_tried(self):
        HarvestedDocument.objects.create(
            url_hash="e" * 64, url="https://x.io/e",
            status=HarvestedDocument.Status.FETCHED, attempts=1, text_chars=5000,
        )
        HarvestedDocument.objects.create(
            url_hash="f" * 64, url="https://x.io/f",
            status=HarvestedDocument.Status.ACCESS_DENIED, attempts=1,
        )
        # Still pending: must not drag the rate down, nothing was asked of it.
        HarvestedDocument.objects.create(url_hash="g" * 64, url="https://x.io/g")

        report = harvest.corpus_report()
        self.assertEqual(report["total"], 3)
        self.assertEqual(report["attempted"], 2)
        self.assertEqual(report["reachable_rate"], 0.5)

    def test_an_empty_corpus_reports_no_rate_rather_than_zero(self):
        self.assertIsNone(harvest.corpus_report()["reachable_rate"])
