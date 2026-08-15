"""Cutting an uploaded document on its own structure instead of on a length.

The archive's chunker (``ExtractionService.text_chunks``) packs sentences up to
a character budget. That is the right rule for a World Bank notice body, which
is one prose block with no headings at all. It is the wrong rule for what a
vendor uploads: a Terms of Reference has sections, a qualification matrix is a
table, and a sentence-packer treats both as a run of sentences that happens to
contain pipes.

Three things follow, and they are the whole module.

**A heading is a boundary and a breadcrumb.** A chunk stops at the next heading
of the same or higher level, and carries the path above it — "3. Qualification
Requirements › 3.2 Financial Capacity" — into its own text. Retrieval then has
the words that place a passage even when the passage itself says only "the
average of the last three years". Without it, the paragraph that states a
threshold and the heading that says what the threshold is *for* are two chunks,
and the search that finds one never sees the other.

**A table is never split.** A row cut away from its header row is a line of
numbers with nothing saying what they are — worse than dropped, because it
still retrieves. Tables are rendered as pipe rows and emitted whole, whatever
the budget says, on the same principle ``ExtractionService`` already applies to
a sentence longer than a chunk.

**This runs for uploads only, and that bound is the point** (D61). Re-chunking
the mirrored archive would change every offset in it and mean re-embedding
74,000 chunks to gain structure that notice bodies do not have. So
``origin=client_supplied`` documents are parsed this way and everything already
indexed keeps the chunker its positions were measured with. The two paths meet
again immediately: both produce ``Chunk`` objects with offsets into a canonical
string that travels with them, and the viewer cannot tell which made which.

**No new dependency.** Markdown is parsed here, ``.docx`` through the
``python-docx`` the harvester already uses, HTML through the standard library's
own parser. A build missing an optional parser yields no blocks and the caller
falls back to the sentence chunker — the same degradation every other reader in
this codebase performs.
"""

from __future__ import annotations

import io
import logging
import re
from dataclasses import dataclass, field
from html.parser import HTMLParser
from typing import Iterable, Sequence

logger = logging.getLogger(__name__)

#: What a block is. ``table`` and ``code`` are the two kinds that are never
#: split; ``text`` is prose and may be packed with its neighbours.
TEXT = "text"
TABLE = "table"
CODE = "code"

#: ATX headings — ``## 3.2 Financial Capacity``. Setext headings (underlined
#: with ``===``) are handled separately below, because a vendor's exported
#: Markdown uses whichever its editor emitted.
_ATX_RE = re.compile(r"^(#{1,6})\s+(.*?)\s*#*\s*$")
_SETEXT_RE = re.compile(r"^(=+|-+)\s*$")
_FENCE_RE = re.compile(r"^\s*(```|~~~)")
#: A Markdown table row: starts and ends with a pipe, or simply contains two.
#: Deliberately loose — exported tables are frequently missing their outer
#: pipes, and a row that is not recognised is a row that gets split away from
#: its header.
_TABLE_ROW_RE = re.compile(r"^\s*\|.*\|\s*$|^[^|\n]*\|[^|\n]*\|")
#: The separator under a table's header row: ``|---|:--:|``.
_TABLE_RULE_RE = re.compile(r"^\s*\|?[\s:|-]{3,}\|?\s*$")

#: The breadcrumb separator written into a chunk's text. A character no
#: borrower's document uses, so a reader can see where the path ends and the
#: passage begins, and a later parser can tell them apart if it has to.
PATH_SEPARATOR = " › "


@dataclass
class Block:
    """One structural unit of a document, with the headings above it."""

    text: str
    kind: str = TEXT
    heading_path: list[str] = field(default_factory=list)
    #: Heading level of the block's own heading, when it *is* one. Zero for
    #: body content. Kept because grouping needs to know that a level-2 chunk
    #: may absorb the level-3 blocks under it but not the next level-2.
    level: int = 0

    @property
    def breadcrumb(self) -> str:
        """The path as it is printed. Padding levels are not printed.

        ``heading_path`` can hold empty entries: a document whose first heading
        is an ``<h2>`` has no level 1, and the path is padded so that *depth*
        stays honest and a later ``<h1>`` still resets everything under it.
        None of that is worth showing a model, and a breadcrumb that opened
        with " › " would read as a heading nobody wrote.
        """
        return PATH_SEPARATOR.join(part for part in self.heading_path if part)


# -- markdown ---------------------------------------------------------------
def blocks_from_markdown(text: str) -> list[Block]:
    """Headings, paragraphs, tables and fenced code, in document order.

    A single pass with three modes, because the alternative — a real Markdown
    parser — would be a dependency for a job whose only outputs are "where does
    this section start" and "is this run of lines a table". Inline formatting
    is left exactly as written: bold markers cost a token and removing them
    would edit a document a vendor is entitled to see quoted verbatim.
    """
    lines = (text or "").splitlines()
    blocks: list[Block] = []
    path: list[str] = []

    buffer: list[str] = []
    mode = TEXT

    def flush() -> None:
        nonlocal buffer, mode
        body = "\n".join(buffer).strip()
        if body:
            blocks.append(Block(text=body, kind=mode, heading_path=list(path)))
        buffer = []
        mode = TEXT

    index = 0
    while index < len(lines):
        line = lines[index]

        if _FENCE_RE.match(line):
            # A fence runs to its closing marker or to the end of the file.
            # Everything between is emitted whole: a code or data block cut in
            # half is a block that parses as neither.
            flush()
            fence = [line]
            index += 1
            while index < len(lines):
                fence.append(lines[index])
                if _FENCE_RE.match(lines[index]):
                    index += 1
                    break
                index += 1
            blocks.append(
                Block(text="\n".join(fence), kind=CODE, heading_path=list(path))
            )
            continue

        heading = _ATX_RE.match(line)
        setext_level = _setext_level(lines, index)
        if heading or setext_level:
            flush()
            if heading:
                level = len(heading.group(1))
                title = heading.group(2).strip()
                index += 1
            else:
                level = setext_level or 1
                title = line.strip()
                # Two lines consumed: the title and the rule under it.
                index += 2
            _descend(path, level, title)
            blocks.append(
                Block(text=title, kind=TEXT, heading_path=list(path), level=level)
            )
            continue

        if _is_table_row(line):
            flush()
            table: list[str] = []
            while index < len(lines) and (
                _is_table_row(lines[index]) or _TABLE_RULE_RE.match(lines[index])
            ):
                table.append(lines[index].rstrip())
                index += 1
            blocks.append(
                Block(text="\n".join(table), kind=TABLE, heading_path=list(path))
            )
            continue

        if not line.strip():
            flush()
            index += 1
            continue

        buffer.append(line.rstrip())
        index += 1

    flush()
    return blocks


def _setext_level(lines: Sequence[str], index: int) -> int:
    """1 or 2 when ``lines[index]`` is a Setext heading, else 0."""
    if index + 1 >= len(lines) or not lines[index].strip():
        return 0
    rule = _SETEXT_RE.match(lines[index + 1])
    if not rule:
        return 0
    # A run of hyphens under a *table row* is the table's rule, not a heading,
    # and a document that opens a section with a dashed line is rarer than a
    # table that does. The table check runs first for that reason.
    if _is_table_row(lines[index]):
        return 0
    return 1 if rule.group(1).startswith("=") else 2


def _descend(path: list[str], level: int, title: str) -> None:
    """Move the breadcrumb to ``level``, replacing what was at or below it.

    Levels skip in real documents — a `#` followed by a `###` — so the path is
    padded rather than indexed. Padding with the heading itself would repeat
    it; padding with an empty string keeps the depth honest and prints nothing.
    """
    del path[level - 1 :]
    while len(path) < level - 1:
        path.append("")
    path.append(title)


def _is_table_row(line: str) -> bool:
    return bool(line.strip()) and bool(_TABLE_ROW_RE.match(line))


# -- docx -------------------------------------------------------------------
def blocks_from_docx(payload: bytes) -> list[Block]:
    """A Word document as headings, paragraphs and pipe tables, in order.

    Body order matters and ``python-docx`` does not hand it over directly:
    ``document.paragraphs`` and ``document.tables`` are two lists, and reading
    them one after the other puts every table at the end of the file — under
    whichever heading happened to be last. So the body's own XML children are
    walked instead, which is the only place the interleaving exists.

    Heading level comes from the paragraph's style name (``Heading 2``), which
    is what a document written in Word actually carries. A document that styles
    its headings by hand yields one flat section, which is the same thing the
    sentence chunker would have produced.
    """
    try:
        import docx  # noqa: PLC0415 - optional, exactly as in the harvester
        from docx.table import Table  # noqa: PLC0415
        from docx.text.paragraph import Paragraph  # noqa: PLC0415
    except ImportError:  # pragma: no cover - packaging guard
        logger.info("python-docx is not installed; no structural chunking for .docx")
        return []

    try:
        document = docx.Document(io.BytesIO(payload))
    except Exception as exc:  # noqa: BLE001 - an unreadable file is a state
        logger.info("Could not read the uploaded .docx: %s", exc)
        return []

    blocks: list[Block] = []
    path: list[str] = []

    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            paragraph = Paragraph(child, document)
            text = (paragraph.text or "").strip()
            if not text:
                continue
            level = _docx_heading_level(paragraph)
            if level:
                _descend(path, level, text)
                blocks.append(
                    Block(text=text, kind=TEXT, heading_path=list(path), level=level)
                )
            else:
                blocks.append(Block(text=text, kind=TEXT, heading_path=list(path)))
        elif tag == "tbl":
            rendered = _docx_table(Table(child, document))
            if rendered:
                blocks.append(
                    Block(text=rendered, kind=TABLE, heading_path=list(path))
                )
    return blocks


def _docx_heading_level(paragraph: object) -> int:
    """The heading level of a paragraph, or 0.

    ``Title`` is level 1 rather than a level of its own: a document has one,
    it sits above everything, and giving it its own tier would push every real
    heading down one and make the breadcrumbs disagree with the numbering the
    document prints.
    """
    name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")
    if name.strip().casefold() == "title":
        return 1
    match = re.match(r"heading\s+(\d)", name.strip().casefold())
    return int(match.group(1)) if match else 0


def _docx_table(table: object) -> str:
    """One Word table as pipe rows, with a rule under the first.

    Rendered rather than flattened because the shape *is* the information: a
    qualification matrix is a grid of criterion against requirement, and a
    reader given its cells in reading order has a list of numbers.
    """
    rows: list[str] = []
    for row in getattr(table, "rows", []):
        cells = [
            " ".join((cell.text or "").split()) for cell in getattr(row, "cells", [])
        ]
        if any(cells):
            rows.append("| " + " | ".join(cells) + " |")
    if not rows:
        return ""
    width = rows[0].count("|") - 1
    rule = "|" + "|".join([" --- "] * max(width, 1)) + "|"
    return "\n".join([rows[0], rule, *rows[1:]])


# -- html -------------------------------------------------------------------
class _TableCollector(HTMLParser):
    """Headings, paragraphs and tables out of HTML, in document order.

    The standard library's parser rather than a dependency, and deliberately
    forgiving: this reads documents a borrower's office software exported, and
    a strict parser's job on those is to raise.
    """

    HEADINGS = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}
    BREAKS = {"p", "div", "li", "br", "tr", "section", "article"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.blocks: list[Block] = []
        self._path: list[str] = []
        self._text: list[str] = []
        self._heading: int = 0
        self._rows: list[list[str]] = []
        self._cells: list[str] | None = None
        self._cell: list[str] | None = None
        self._in_table = 0

    # -- assembly -------------------------------------------------------
    def _flush_text(self) -> None:
        body = " ".join(" ".join(self._text).split())
        self._text = []
        if not body:
            return
        if self._heading:
            _descend(self._path, self._heading, body)
            self.blocks.append(
                Block(
                    text=body,
                    kind=TEXT,
                    heading_path=list(self._path),
                    level=self._heading,
                )
            )
        else:
            self.blocks.append(
                Block(text=body, kind=TEXT, heading_path=list(self._path))
            )

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag == "table":
            self._flush_text()
            self._in_table += 1
            self._rows = []
            return
        if self._in_table:
            if tag == "tr":
                self._cells = []
            elif tag in {"td", "th"}:
                self._cell = []
            return
        if tag in self.HEADINGS:
            self._flush_text()
            self._heading = self.HEADINGS[tag]
        elif tag in self.BREAKS:
            self._flush_text()

    def handle_endtag(self, tag: str) -> None:
        if self._in_table:
            if tag in {"td", "th"} and self._cells is not None:
                self._cells.append(" ".join("".join(self._cell or []).split()))
                self._cell = None
            elif tag == "tr" and self._cells is not None:
                if any(self._cells):
                    self._rows.append(self._cells)
                self._cells = None
            elif tag == "table":
                self._in_table = max(self._in_table - 1, 0)
                rendered = _rows_to_pipes(self._rows)
                self._rows = []
                if rendered and not self._in_table:
                    self.blocks.append(
                        Block(
                            text=rendered, kind=TABLE, heading_path=list(self._path)
                        )
                    )
            return
        if tag in self.HEADINGS:
            self._flush_text()
            self._heading = 0
        elif tag in self.BREAKS:
            self._flush_text()

    def handle_data(self, data: str) -> None:
        if self._cell is not None:
            self._cell.append(data)
        elif not self._in_table:
            self._text.append(data)

    def close(self) -> None:  # noqa: D102 - inherited contract
        super().close()
        self._flush_text()


def blocks_from_html(markup: str) -> list[Block]:
    """HTML as blocks, with every ``<table>`` kept as a grid."""
    parser = _TableCollector()
    try:
        parser.feed(markup or "")
        parser.close()
    except Exception as exc:  # noqa: BLE001 - malformed markup is a state
        logger.info("Could not parse the uploaded HTML: %s", exc)
        return parser.blocks
    return parser.blocks


def _rows_to_pipes(rows: Sequence[Sequence[str]]) -> str:
    if not rows:
        return ""
    lines = ["| " + " | ".join(row) + " |" for row in rows if any(row)]
    if not lines:
        return ""
    width = max(len(row) for row in rows)
    rule = "|" + "|".join([" --- "] * max(width, 1)) + "|"
    return "\n".join([lines[0], rule, *lines[1:]])


# -- rendering and grouping -------------------------------------------------
def render(blocks: Sequence[Block]) -> str:
    """The document as one string, in the order the blocks were found.

    This is the text everything downstream measures against: the chunker finds
    its offsets in it, and the viewer serves it beside them. So it is built
    once, here, from the same block list the chunks come from — the failure
    this avoids is the one ``chunks.py`` opens by naming, a highlight drawn
    from one parse over a page rendered from another.
    """
    return "\n\n".join(block.text for block in blocks if block.text.strip())


def group(
    blocks: Sequence[Block], *, chunk_chars: int, min_chunk_chars: int
) -> Iterable[list[Block]]:
    """Consecutive blocks packed into chunks that respect the structure.

    Three rules, in order of precedence:

    * a heading starts a new chunk, and the heading goes *with* what follows it
      rather than closing what came before;
    * a table or a code block is its own chunk, whatever the budget says;
    * everything else packs up to ``chunk_chars``.

    ``min_chunk_chars`` is not applied here. A section whose whole body is one
    short line is still that section, and dropping it would lose a heading the
    breadcrumb of every following chunk no longer mentions. The caller drops
    what is too small to embed, which is where that judgement already lives.
    """
    batch: list[Block] = []
    size = 0

    for block in blocks:
        atomic = block.kind in {TABLE, CODE}
        starts_section = block.level > 0

        if batch and (atomic or starts_section or size + len(block.text) > chunk_chars):
            yield batch
            batch = []
            size = 0

        batch.append(block)
        size += len(block.text)

        if atomic:
            yield batch
            batch = []
            size = 0

    if batch:
        yield batch


def chunk_text(batch: Sequence[Block]) -> str:
    """One chunk's text: its breadcrumb, then its blocks.

    The breadcrumb is prepended rather than stored beside the passage because
    it has to reach the *embedding*. A path kept in the payload would help a
    reader and do nothing for retrieval, and retrieval is the reason a chunk
    that says only "the average of the last three years" needs to know it sits
    under "Financial Capacity".

    Printed once per chunk, not once per block: the blocks in a batch share a
    path by construction, and repeating it would weight the section title
    above everything the section says.
    """
    if not batch:
        return ""
    body = "\n\n".join(block.text for block in batch if block.text.strip())
    head = batch[0].breadcrumb
    if not head:
        return body
    # A chunk that *opens* with its own heading would otherwise print it twice:
    # once as the breadcrumb's last element and once as the block itself.
    first = batch[0].text.strip()
    if batch[0].level and head.endswith(first):
        return f"{head}\n{body[len(first):].lstrip()}".strip()
    return f"{head}\n{body}"
